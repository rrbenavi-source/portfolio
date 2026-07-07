#!/usr/bin/env python3
"""Genera un CV en Word (.docx) editable para Ricardo Benavides."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0F, 0x7A, 0x5A)   # verde SAP/dato sobrio
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Márgenes
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)

# Fuente base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(4)


def hr(p):
    """Línea inferior fina en un párrafo (separador de sección)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '0F7A5A')
    pbdr.append(bottom); pPr.append(pbdr)


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    hr(p)
    return p


def bullet(text_runs):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for txt, bold in text_runs:
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(10)
    return p


def role(title, company, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK
    if company:
        r2 = p.add_run('  —  ' + company); r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = ACCENT
    # tabulador derecho para la fecha
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.1), WD_ALIGN_PARAGRAPH.RIGHT)
    r3 = p.add_run('\t' + date); r3.italic = True; r3.font.size = Pt(9); r3.font.color.rgb = GREY
    return p


# ---------- Encabezado ----------
name = doc.add_paragraph()
name.paragraph_format.space_after = Pt(0)
rn = name.add_run('RICARDO BENAVIDES')
rn.bold = True; rn.font.size = Pt(22); rn.font.color.rgb = DARK; rn.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
rs = sub.add_run('Líder de BI & Arquitecto de Datos SAP  ·  Integración SAP & Databricks')
rs.font.size = Pt(11.5); rs.font.color.rgb = ACCENT; rs.bold = True

contact = doc.add_paragraph()
rc = contact.add_run('Monterrey, Nuevo León, México   |   rbenavides@dcsol.com.mx   |   rrbenavi@gmail.com   |   linkedin.com/in/rrbenavi')
rc.font.size = Pt(9.5); rc.font.color.rgb = GREY
hr(contact)

# ---------- Perfil ----------
heading('Perfil profesional')
p = doc.add_paragraph()
for txt, b in [
    ('Líder en Business Intelligence y Arquitectura de Datos SAP con ', False),
    ('más de 19 años de experiencia', True),
    (' diseñando, implementando y optimizando soluciones de datos que impulsan decisiones estratégicas. '
     'Especialista en ', False),
    ('SAP BW/4HANA, SAP Analytics Cloud (SAC), SAP BPC y SAP BusinessObjects', True),
    (', con integración SAP–Databricks. Como líder, dirijo equipos multidisciplinarios y proyectos de alto '
     'impacto, combinando visión de negocio y profundidad técnica para convertir datos en ventajas '
     'competitivas medibles en industria, consumo masivo, sector hotelero y gobierno.', False),
]:
    r = p.add_run(txt); r.bold = b; r.font.size = Pt(10)

bullet([('Reducción de tiempos de análisis en más de ', False), ('40%', True), ('.', False)])
bullet([('Mejora de la precisión de indicadores clave para ', False), ('12+ áreas de negocio', True), ('.', False)])
bullet([('Beneficio económico de ', False), ('€1.5M', True), (' generado en una solución de inteligencia de datos para manufactura.', False)])
bullet([('50%', True), (' de ahorro en licenciamiento de SAP Analytics Cloud mediante estrategia de migración.', False)])

# ---------- Competencias ----------
heading('Competencias clave')
def comp(label, body):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + ': '); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
    r2 = p.add_run(body); r2.font.size = Pt(10)
comp('Arquitectura & Plataformas',
     'SAP BW/4HANA · SAP Analytics Cloud (SAC) · SAP BPC · SAP BusinessObjects (Web Intelligence, Analysis OLAP, '
     'Design Studio, Lumira) · SAP Data Intelligence · Analysis for Office · Microsoft Azure · Databricks · ABAP')
comp('Disciplinas',
     'Arquitectura y modelado de datos · Analítica avanzada e inteligencia prescriptiva · Gobierno de datos · '
     'Data Volume Management · Optimización de desempeño · Tableros ejecutivos · Estrategia de datos y modernización')
comp('Liderazgo',
     'Dirección de equipos multidisciplinarios · Gestión de proyectos · Diseño de propuestas técnicas · '
     'Optimización de costos y licenciamiento · Alineación negocio–TI')

# ---------- Experiencia ----------
heading('Experiencia profesional')

role('Arquitecto SAP BW/4HANA · SAP Analytics Cloud', 'HEINEKEN México', 'abr. 2023 – Actualidad')
doc.add_paragraph('Líder del equipo de Business Intelligence SAP y arquitecto de la plataforma de datos; '
                  'adicionalmente impulso la integración SAP–Databricks.').runs[0].font.size = Pt(9.5)
bullet([('Migración SAP BO → SAP SAC: ', True), ('diseño, liderazgo y ejecución; propuesta de optimización de licencias de planning con ahorro de licenciamiento.', False)])
bullet([('Estrategia SAC → Analysis for Office: ', True), ('reducción del 50% del costo de licenciamiento de SAP Analytics Cloud (2024 y 2025).', False)])
bullet([('Data Volume Management: ', True), ('−50% de espacio en base de datos y −33% en servidores aplicativos.', False)])
bullet([('Migraciones SAP BW/4HANA 2023 y SAP Data Intelligence ', True), ('ejecutadas sin contratiempos.', False)])
bullet([('Alta disponibilidad: ', True), ('resolución del incidente crítico (dic. 2025) de ventas, distribución y pedidos por falla en el sistema fuente.', False)])
bullet([('Automatización ', True), ('del seguimiento diario de ventas a nivel OpCo.', False)])

role('Líder Técnico — SAP Analytics Cloud', 'HEINEKEN México', '2023')
bullet([('Líder de equipo de 12 recursos en la migración de 300+ reportes de SAP BO a SAP Analytics Cloud: ', True),
        ('estrategia, templates, coordinación, resolución técnica y puesta en productivo.', False)])

role('Consultor SAP BW / BO / BPC (Sr.)', 'HEINEKEN México', 'ago. 2019 – abr. 2023')
bullet([('Soluciones BI: W&T2 (logística), Commercial Scorecard, Cartera Nacional, CoVa, P&CI, SD–Lista de Precios.', False)])
bullet([('Migración SAP BPC 10.1 → 11.1: post-instalación, respaldo/restauración del modelo de consolidación y remediación.', False)])

role('Consultor SAP BW (Sr.)', 'Grupo AlEn', 'ago. 2018 – ago. 2019')
bullet([('Soluciones BW para indicadores Nielsen (Scantrack), Foto del Éxito (rentabilidad) y Recursos Humanos sobre BW on HANA.', False)])

role('Líder Técnico SAP BI/BPC · SAP Project Manager', 'Grupo Pueblo Bonito', 'nov. 2016 – ago. 2018')
bullet([('Consolidación financiera del grupo, reportes de operación hotelera (Estado de Resultados), IFRS y solución Mobile.', False)])
bullet([('Gestión de proyectos: BPC, estabilización FI/CO/MM, Portal de Proveedores y Contabilidad Electrónica.', False)])

role('SAP BPC Technical Lead · Senior SAP BW Consultant', 'Fibra Inn', 'nov. 2015 – ene. 2017')
bullet([('Planeación de presupuesto de operaciones hoteleras; desarrollo BI para módulos financieros, reportes BMV y Mobile.', False)])

role('Senior SAP BW Consultant', 'CEMEX', 'feb. 2014 – oct. 2015')
bullet([('Modelos de rentabilidad (México y Colombia), cartera legal/vencida e indicadores de RH integrando fuentes SAP y no SAP.', False)])

# ---------- Caso de éxito ----------
heading('Caso de éxito destacado — D&C Solutions')
p = doc.add_paragraph()
for txt, b in [
    ('Solución integral de inteligencia de datos para un cliente del sector manufacturero con múltiples plantas, '
     'que unifica y estandariza la gestión de contratos de suministro, pasando de un control manual y reactivo a uno '
     'proactivo. Impacto: ', False),
    ('beneficio económico proyectado de €1.5M', True),
    (' (€1.4M en ahorros directos + €100k en costos evitados), ', False),
    ('sin inversión adicional en infraestructura tecnológica', True), ('.', False),
]:
    r = p.add_run(txt); r.bold = b; r.font.size = Pt(10)

# ---------- Clientes ----------
heading('Clientes y organizaciones')
p = doc.add_paragraph()
r = p.add_run('CEMEX · NEMAK · VITRO · Grupo AlEn · Fibra Inn · Grupo Pueblo Bonito · Papel San Francisco · SADM · Mega Alimentos · Peñafiel · HEINEKEN México · Minera México')
r.font.size = Pt(10)

# ---------- Formación / Certificación / Idiomas ----------
heading('Formación · Certificaciones · Idiomas')
bullet([('Licenciado en Ciencias Computacionales (TI)', True), (' — Universidad Autónoma de Nuevo León, Facultad de Ciencias Físico Matemáticas · 2002–2006.', False)])
bullet([('SAP NetWeaver 2004s — Business Intelligence', True), (' · SAP · oct. 2008 · ID 0006064683.', False)])
bullet([('Idiomas: ', True), ('Español (nativo) · Inglés (intermedio).', False)])

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'cv-ricardo-benavides.docx')
doc.save(out)
print('OK ->', out)
